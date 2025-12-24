"""OCR Service for document text extraction."""

import io
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import cv2
import numpy as np
from pdf2image import convert_from_bytes
from PIL import Image
from PyPDF2 import PdfReader

from app.config import settings
from app.utils.logger import get_logger

logger = get_logger(__name__)


class OCRService:
    """OCR service supporting multiple engines."""

    def __init__(self):
        """Initialize OCR service."""
        self.tesseract_available = self._check_tesseract()
        self.paddleocr_available = self._check_paddleocr()
        self.paddle_ocr = None
        
        if self.paddleocr_available:
            try:
                from paddleocr import PaddleOCR
                self.paddle_ocr = PaddleOCR(
                    use_angle_cls=True,
                    lang=settings.PADDLEOCR_LANG,
                    use_gpu=settings.PADDLEOCR_USE_GPU,
                    show_log=False,
                )
                logger.info("PaddleOCR initialized successfully")
            except Exception as e:
                logger.warning("Failed to initialize PaddleOCR", error=str(e))
                self.paddleocr_available = False

    def _check_tesseract(self) -> bool:
        """Check if Tesseract is available."""
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            logger.info("Tesseract OCR available")
            return True
        except Exception as e:
            logger.warning("Tesseract not available", error=str(e))
            return False

    def _check_paddleocr(self) -> bool:
        """Check if PaddleOCR is available."""
        try:
            import paddleocr
            return True
        except ImportError:
            return False

    def extract_text_from_pdf(
        self, pdf_content: bytes, use_ocr: bool = True
    ) -> Dict[str, any]:
        """
        Extract text from PDF.
        
        Args:
            pdf_content: PDF file content as bytes
            use_ocr: Whether to use OCR for scanned PDFs
            
        Returns:
            Dictionary with extracted text, page count, and metadata
        """
        logger.info("Extracting text from PDF")
        
        # Try digital text extraction first
        try:
            text_content = self._extract_digital_pdf_text(pdf_content)
            
            # If we got meaningful text, return it
            if text_content and len(text_content.strip()) > 100:
                return {
                    "text": text_content,
                    "method": "digital",
                    "page_count": len(PdfReader(io.BytesIO(pdf_content)).pages),
                    "is_scanned": False,
                }
        except Exception as e:
            logger.warning("Digital PDF extraction failed", error=str(e))
        
        # Fall back to OCR
        if use_ocr:
            logger.info("Using OCR for PDF extraction")
            return self._extract_pdf_with_ocr(pdf_content)
        
        raise ValueError("Could not extract text from PDF")

    def _extract_digital_pdf_text(self, pdf_content: bytes) -> str:
        """Extract text from digital PDF."""
        reader = PdfReader(io.BytesIO(pdf_content))
        
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num + 1}", error=str(e))
        
        return "\n\n".join(text_parts)

    def _extract_pdf_with_ocr(self, pdf_content: bytes) -> Dict[str, any]:
        """Extract text from PDF using OCR."""
        # Convert PDF to images
        images = convert_from_bytes(pdf_content, dpi=300)
        
        text_parts = []
        for page_num, image in enumerate(images):
            logger.info(f"Processing page {page_num + 1}/{len(images)} with OCR")
            
            # Preprocess image
            preprocessed = self._preprocess_image(image)
            
            # Extract text
            page_text = self._ocr_image(preprocessed)
            
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        return {
            "text": "\n\n".join(text_parts),
            "method": "ocr",
            "page_count": len(images),
            "is_scanned": True,
        }

    def extract_text_from_image(self, image_content: bytes) -> Dict[str, any]:
        """
        Extract text from image.
        
        Args:
            image_content: Image file content as bytes
            
        Returns:
            Dictionary with extracted text and metadata
        """
        logger.info("Extracting text from image")
        
        # Load image
        image = Image.open(io.BytesIO(image_content))
        
        # Preprocess
        preprocessed = self._preprocess_image(image)
        
        # Extract text
        text = self._ocr_image(preprocessed)
        
        return {
            "text": text,
            "method": "ocr",
            "page_count": 1,
            "is_scanned": True,
        }

    def extract_text_from_docx(self, docx_content: bytes) -> Dict[str, any]:
        """Extract text from DOCX file."""
        from docx import Document
        
        logger.info("Extracting text from DOCX")
        
        doc = Document(io.BytesIO(docx_content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return {
            "text": "\n\n".join(text_parts),
            "method": "digital",
            "page_count": len(doc.sections),
            "is_scanned": False,
        }

    def _preprocess_image(self, image: Image.Image) -> np.ndarray:
        """
        Preprocess image for better OCR results.
        
        Applies:
        - Grayscale conversion
        - Noise reduction
        - Thresholding
        - Deskewing
        """
        # Convert PIL Image to OpenCV format
        img_array = np.array(image)
        
        # Convert to grayscale
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Apply adaptive thresholding
        thresh = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        
        # Deskew
        deskewed = self._deskew_image(thresh)
        
        return deskewed

    def _deskew_image(self, image: np.ndarray) -> np.ndarray:
        """Deskew an image."""
        # Find all white pixels
        coords = np.column_stack(np.where(image > 0))
        
        # Find minimum area rectangle
        angle = cv2.minAreaRect(coords)[-1]
        
        # Adjust angle
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        
        # Rotate image
        (h, w) = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(
            image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
        )
        
        return rotated

    def _ocr_image(self, image: np.ndarray) -> str:
        """
        Perform OCR on preprocessed image.
        
        Uses PaddleOCR if available, falls back to Tesseract.
        """
        # Try PaddleOCR first (better for financial documents)
        if self.paddleocr_available and self.paddle_ocr:
            try:
                result = self.paddle_ocr.ocr(image, cls=True)
                
                # Extract text from result
                text_parts = []
                if result and result[0]:
                    for line in result[0]:
                        if line and len(line) > 1:
                            text_parts.append(line[1][0])  # Get text from result
                
                return "\n".join(text_parts)
            except Exception as e:
                logger.warning("PaddleOCR failed, falling back to Tesseract", error=str(e))
        
        # Fall back to Tesseract
        if self.tesseract_available:
            try:
                import pytesseract
                
                # Convert numpy array to PIL Image
                pil_image = Image.fromarray(image)
                
                text = pytesseract.image_to_string(
                    pil_image,
                    lang=settings.TESSERACT_LANG,
                    config='--psm 3'  # Fully automatic page segmentation
                )
                
                return text
            except Exception as e:
                logger.error("Tesseract OCR failed", error=str(e), exc_info=True)
                raise
        
        raise RuntimeError("No OCR engine available")

    def extract_tables(self, image: np.ndarray) -> List[List[List[str]]]:
        """
        Extract tables from image.
        
        Returns list of tables, where each table is a list of rows,
        and each row is a list of cell contents.
        """
        # This is a simplified version - for production, use libraries like
        # table-transformer or camelot-py
        logger.info("Table extraction not fully implemented - returning empty list")
        return []


# Global OCR service instance
ocr_service = OCRService()


def get_ocr_service() -> OCRService:
    """Get OCR service instance."""
    return ocr_service
